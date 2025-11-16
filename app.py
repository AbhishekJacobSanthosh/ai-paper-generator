from flask import Flask, request, jsonify, render_template, send_file
import requests
import json
from datetime import datetime
import os
import uuid
import re
from io import BytesIO
import concurrent.futures
import hashlib
import base64
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Frame, PageTemplate, BaseDocTemplate, Table, TableStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib import colors
from reportlab.platypus import Image as RLImage
from PIL import Image, ImageEnhance
import easyocr

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs('saved_papers', exist_ok=True)

print("Loading EasyOCR model...")
reader = easyocr.Reader(['en'], gpu=False)
print("EasyOCR loaded!")

OLLAMA_API_URL = "http://localhost:11434/api/generate"
MODEL_NAME = "mistral"
WINSTON_API_KEY = "tpctrVgy31ssAG8iFX5kZlNITqzfAwfFAviPtzBT34b2c160" 

CITATION_DATABASE = {
    "machine learning": [
        "[1] I. Goodfellow, Y. Bengio, and A. Courville, Deep Learning. MIT Press, 2016.",
        "[2] C. Bishop, Pattern Recognition and Machine Learning. Springer, 2006.",
        "[3] T. Mitchell, Machine Learning. McGraw-Hill, 1997.",
        "[4] Y. LeCun, Y. Bengio, and G. Hinton, \"Deep learning,\" Nature, vol. 521, pp. 436-444, 2015.",
        "[5] A. Ng, \"Machine Learning Yearning,\" Technical Strategy, 2018."
    ],
    "neural networks": [
        "[1] G. Hinton et al., \"A fast learning algorithm for deep belief nets,\" Neural Comp., vol. 18, 2006.",
        "[2] A. Krizhevsky et al., \"ImageNet classification with deep CNNs,\" in Proc. NIPS, 2012."
    ],
    "default": [
        "[1] A. Author, B. Author, \"Title of Paper,\" Journal Name, vol. X, no. Y, pp. ZZ-ZZ, 2023.",
        "[2] C. Author et al., \"Conference Paper Title,\" in Proc. Conf. Name, 2024, pp. XX-YY.",
        "[3] D. Author, \"Book Title,\" Publisher, 2022.",
        "[4] E. Author and F. Author, \"Another Paper,\" IEEE Trans., vol. 10, 2023.",
        "[5] G. Author, \"Recent Work,\" ACM Computing Surveys, 2024."
    ]
}

def generate_text(prompt, max_tokens=400, temperature=0.7):
    try:
        response = requests.post(OLLAMA_API_URL, json={
            "model": MODEL_NAME,
            "prompt": prompt,
            "stream": False,
            "options": {
                "num_predict": max_tokens,
                "temperature": temperature,
                "top_p": 0.9,
                "top_k": 40
            }
        }, timeout=150)
        
        if response.status_code == 200:
            result = response.json()
            return result.get('response', '').strip()
        return "Error: Could not generate text"
    except Exception as e:
        return f"Error: {str(e)}"

def generate_section(section_name, prompt, max_tokens=400):
    try:
        result = generate_text(prompt, max_tokens=max_tokens)
        return section_name, result if not result.startswith("Error") else "[Content generation failed]"
    except:
        return section_name, "[Error]"

def extract_text_from_image(image_path):
    try:
        if not os.path.exists(image_path):
            return "OCR Error: Image not found"
        
        img = Image.open(image_path)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.5)
        
        results = reader.readtext(image_path, detail=0, paragraph=False)
        text = ' '.join(results).strip()
        
        return text if text else "No text detected"
    except Exception as e:
        return f"OCR Error: {str(e)}"

def check_plagiarism(text):
    """Enhanced plagiarism check with Winston AI and local fallback"""
    
    # Try Winston AI first if API key is available
    if WINSTON_API_KEY:
        try:
            print("🔍 Checking plagiarism with Winston AI...")
            
            # Limit to first 5000 words to save credits
            text_to_check = ' '.join(text.split()[:5000])
            
            response = requests.post(
                "https://api.gowinston.ai/v1/plagiarism",
                headers={
                    "Authorization": f"Bearer {WINSTON_API_KEY}",
                    "Content-Type": "application/json"
                },
                json={
                    "text": text_to_check,
                    "language": "en"
                },
                timeout=60
            )
            
            if response.status_code == 200:
                data = response.json()
                plagiarism_pct = data.get("plagiarism_percentage", 0)
                
                print(f"✓ Winston AI: {100 - plagiarism_pct}% unique")
                
                return {
                    "uniqueness_score": round(100 - plagiarism_pct, 1),
                    "plagiarism_percentage": plagiarism_pct,
                    "status": "Highly Original" if plagiarism_pct < 5 else 
                             "Original" if plagiarism_pct < 15 else 
                             "Needs Review",
                    "sources_found": len(data.get("sources", [])),
                    "total_sentences": len(text.split('.')),
                    "api_provider": "Winston AI ✓",
                    "is_authentic": True
                }
            else:
                print(f"⚠ Winston API error: {response.status_code}")
                
        except requests.exceptions.Timeout:
            print("⚠ Winston API timeout")
        except Exception as e:
            print(f"⚠ Winston API error: {e}")
    
    # Fallback to enhanced local checking
    print("🔍 Using local plagiarism analysis...")
    
    sentences = [s.strip() for s in text.split('.') if s.strip()]
    unique_sentences = len(set([s.lower() for s in sentences]))
    uniqueness = (unique_sentences / len(sentences) * 100) if sentences else 0
    
    # Vocabulary diversity
    words = re.findall(r'\w+', text.lower())
    unique_words = len(set(words))
    vocabulary_score = (unique_words / len(words) * 100) if words else 0
    
    # Common academic phrases detection
    common_phrases = [
        "in conclusion", "as a result", "however", "furthermore",
        "this study", "research shows", "according to", "it is important",
        "in addition", "on the other hand"
    ]
    common_count = sum(1 for phrase in common_phrases if phrase in text.lower())
    
    # Sentence length variance (AI text often uniform)
    sentence_lengths = [len(s.split()) for s in sentences if s]
    if sentence_lengths:
        avg_length = sum(sentence_lengths) / len(sentence_lengths)
        variance = sum((l - avg_length) ** 2 for l in sentence_lengths) / len(sentence_lengths)
    else:
        variance = 0
    
    # Calculate weighted score
    final_score = (
        uniqueness * 0.4 +
        min(vocabulary_score, 100) * 0.3 +
        min(100 - (common_count * 3), 100) * 0.2 +
        min(variance * 1.5, 100) * 0.1
    )
    
    return {
        "uniqueness_score": round(final_score, 1),
        "status": "Highly Original" if final_score > 85 else "Original" if final_score > 70 else "Needs Review",
        "total_sentences": len(sentences),
        "unique_sentences": unique_sentences,
        "vocabulary_diversity": round(vocabulary_score, 1),
        "common_phrases_detected": common_count,
        "avg_sentence_length": round(avg_length if sentence_lengths else 0, 1),
        "api_provider": "Local Analysis (Enhanced)",
        "is_authentic": False,
        "note": "For production use, add Winston API key for real web-based checking"
    }

def get_citations(topic):
    topic_lower = topic.lower()
    for key in CITATION_DATABASE.keys():
        if key in topic_lower:
            return CITATION_DATABASE[key]
    return CITATION_DATABASE["default"]

def generate_doi():
    return f"10.1109/ACCESS.{datetime.now().year}.{uuid.uuid4().hex[:8].upper()}"

def generate_sample_bar_chart(topic):
    plt.figure(figsize=(6, 4))
    categories = ['Method A', 'Method B', 'Method C', 'Method D', 'Proposed']
    values = [72, 78, 81, 85, 92]
    colors_list = ['#667eea', '#764ba2', '#f093fb', '#f5576c', '#28a745']
    
    plt.bar(categories, values, color=colors_list, alpha=0.8, edgecolor='black', linewidth=0.7)
    plt.xlabel('Approaches', fontsize=11, fontweight='bold')
    plt.ylabel('Accuracy (%)', fontsize=11, fontweight='bold')
    plt.title(f'Performance Comparison', fontsize=12, fontweight='bold')
    plt.ylim(0, 100)
    plt.grid(axis='y', alpha=0.3, linestyle='--')
    
    for i, v in enumerate(values):
        plt.text(i, v + 2, f'{v}%', ha='center', fontweight='bold', fontsize=9)
    
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    return buf

def generate_sample_line_chart(topic):
    plt.figure(figsize=(6, 4))
    epochs = np.arange(1, 11)
    train_acc = 65 + 20 * (1 - np.exp(-epochs/3)) + np.random.uniform(-2, 2, 10)
    val_acc = 60 + 18 * (1 - np.exp(-epochs/3)) + np.random.uniform(-2, 2, 10)
    
    plt.plot(epochs, train_acc, marker='o', linewidth=2, label='Training', color='#667eea', markersize=6)
    plt.plot(epochs, val_acc, marker='s', linewidth=2, label='Validation', color='#f5576c', markersize=6)
    
    plt.xlabel('Epoch', fontsize=11, fontweight='bold')
    plt.ylabel('Accuracy (%)', fontsize=11, fontweight='bold')
    plt.title(f'Training Progress', fontsize=12, fontweight='bold')
    plt.legend(loc='lower right', fontsize=10)
    plt.grid(True, alpha=0.3, linestyle='--')
    plt.ylim(50, 95)
    
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    return buf

def create_sample_table_data(topic):
    return [
        ['Method', 'Accuracy', 'Precision', 'Recall', 'F1-Score'],
        ['Baseline', '72.3%', '71.5%', '70.8%', '71.1%'],
        ['Method A', '78.6%', '77.2%', '79.1%', '78.1%'],
        ['Method B', '81.4%', '80.8%', '82.3%', '81.5%'],
        ['Proposed', '92.1%', '91.8%', '92.4%', '92.1%']
    ]

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/api/generate-paper', methods=['POST'])
def generate_paper():
    data = request.json
    topic = data.get('topic', '')
    author_name = data.get('author_name', 'Author Name')
    affiliation = data.get('affiliation', 'University Name')
    email = data.get('email', 'author@university.edu')
    
    if not topic:
        return jsonify({"success": False, "error": "Topic required"}), 400
    
    paper = {
        "title": topic,
        "author": author_name,
        "affiliation": affiliation,
        "email": email,
        "doi": generate_doi(),
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "sections": {},
        "metadata": {},
        "figures": {}
    }
    
    prompts = {
        "abstract": (f"""Write a 180-word abstract for: "{topic}". Include: background, objective, methods, results, conclusion.""", 280),
        "introduction": (f"""Write 350-word Introduction for: "{topic}". Cover: context, problem, objectives, structure.""", 450),
        "literature_review": (f"""Write 350-word Literature Review for: "{topic}". Discuss: key research, gaps.""", 450),
        "methodology": (f"""Write 350-word Methodology for: "{topic}". Include: design, data collection, analysis.""", 450),
        "results": (f"""Write 300-word Results for: "{topic}". Mention Figure 1 shows comparison and Table 1 presents metrics.""", 400),
        "discussion": (f"""Write 350-word Discussion for: "{topic}". Cover: interpretation, implications, limitations.""", 450),
        "conclusion": (f"""Write 250-word Conclusion for: "{topic}". Summarize findings and future work.""", 350),
    }
    
    print(f"⚡ Starting generation: {topic}")
    
    with concurrent.futures.ThreadPoolExecutor(max_workers=6) as executor:
        future_to_section = {
            executor.submit(generate_section, key, prompt, tokens): key 
            for key, (prompt, tokens) in prompts.items()
        }
        
        for future in concurrent.futures.as_completed(future_to_section):
            section_key, result = future.result()
            paper["sections"][section_key] = result
            print(f"✓ {section_key}")
    
    print("📊 Generating figures...")
    
    try:
        fig1_buf = generate_sample_bar_chart(topic)
        fig1_base64 = base64.b64encode(fig1_buf.getvalue()).decode('utf-8')
        paper["figures"]["figure1"] = {
            "caption": "Figure 1: Performance comparison of different approaches",
            "data": fig1_base64,
            "type": "bar_chart"
        }
        
        fig2_buf = generate_sample_line_chart(topic)
        fig2_base64 = base64.b64encode(fig2_buf.getvalue()).decode('utf-8')
        paper["figures"]["figure2"] = {
            "caption": "Figure 2: Training and validation accuracy over epochs",
            "data": fig2_base64,
            "type": "line_chart"
        }
        
        paper["figures"]["table1"] = {
            "caption": "Table 1: Comparative performance metrics",
            "data": create_sample_table_data(topic),
            "type": "table"
        }
        
        print("✓ Figures generated")
    except Exception as e:
        print(f"Warning: {e}")
    
    paper["sections"]["references"] = "\n".join(get_citations(topic))
    
    total_words = sum(len(content.split()) for content in paper["sections"].values())
    paper["metadata"] = {
        "word_count": total_words,
        "page_estimate": round(total_words / 250),
        "has_figures": len(paper.get("figures", {}))
    }
    
    print(f"✅ Done: {total_words} words, {len(paper['figures'])} figures")
    
    return jsonify({"success": True, "paper": paper})

@app.route('/api/plagiarism-check', methods=['POST'])
def plagiarism_check():
    data = request.json
    paper = data.get('paper', {})
    
    if not paper:
        return jsonify({"success": False, "error": "No paper"}), 400
    
    full_text = " ".join(paper.get('sections', {}).values())
    result = check_plagiarism(full_text)
    
    return jsonify({"success": True, "plagiarism_check": result})

@app.route('/api/ocr-generate', methods=['POST'])
def ocr_generate():
    if 'image' not in request.files:
        return jsonify({"success": False, "error": "No image"}), 400
    
    file = request.files['image']
    if not file.filename:
        return jsonify({"success": False, "error": "Empty filename"}), 400
    
    try:
        file_ext = os.path.splitext(file.filename)[1].lower() or '.jpg'
        safe_filename = f"ocr_{uuid.uuid4().hex[:8]}{file_ext}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)
        
        file.save(filepath)
        extracted_text = extract_text_from_image(filepath)
        
        try:
            os.remove(filepath)
        except:
            pass
        
        return jsonify({"success": True, "extracted_text": extracted_text})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/export-pdf', methods=['POST'])
def export_pdf():
    data = request.json
    paper = data.get('paper', {})
    
    if not paper:
        return jsonify({"success": False, "error": "No paper"}), 400
    
    buffer = BytesIO()
    doc = BaseDocTemplate(buffer, pagesize=letter,
                         leftMargin=0.75*inch, rightMargin=0.75*inch,
                         topMargin=0.75*inch, bottomMargin=0.75*inch)
    
    frame_width = (letter[0] - 1.5*inch - 0.2*inch) / 2
    frame_height = letter[1] - 1.5*inch
    
    left_frame = Frame(0.75*inch, 0.75*inch, frame_width, frame_height, id='left')
    right_frame = Frame(0.75*inch + frame_width + 0.2*inch, 0.75*inch, frame_width, frame_height, id='right')
    
    doc.addPageTemplates([PageTemplate(id='TwoCol', frames=[left_frame, right_frame])])
    
    elements = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, leading=22,
                                 alignment=TA_CENTER, fontName='Times-Bold', spaceAfter=12)
    author_style = ParagraphStyle('Author', parent=styles['Normal'], fontSize=10, leading=12,
                                   alignment=TA_CENTER, fontName='Times-Roman', spaceAfter=3)
    section_style = ParagraphStyle('Section', parent=styles['Heading2'], fontSize=10, leading=12,
                                    fontName='Times-Bold', spaceAfter=6, spaceBefore=8)
    body_style = ParagraphStyle('Body', parent=styles['BodyText'], fontSize=9, leading=11,
                                alignment=TA_JUSTIFY, fontName='Times-Roman', firstLineIndent=12, spaceAfter=6)
    abstract_style = ParagraphStyle('Abstract', parent=body_style, fontName='Times-Italic', firstLineIndent=0)
    caption_style = ParagraphStyle('Caption', parent=body_style, fontSize=8, alignment=TA_CENTER, 
                                   fontName='Times-Italic', firstLineIndent=0)
    
    elements.append(Paragraph(paper.get('title', 'Research Paper'), title_style))
    elements.append(Spacer(1, 0.1*inch))
    
    author_info = f"{paper.get('author', 'Author')}<br/>{paper.get('affiliation', 'Affiliation')}<br/>Email: {paper.get('email', 'email@edu')}"
    elements.append(Paragraph(author_info, author_style))
    elements.append(Spacer(1, 0.15*inch))
    
    if 'abstract' in paper.get('sections', {}):
        elements.append(Paragraph("<b><i>Abstract</i></b>—", section_style))
        elements.append(Paragraph(paper['sections']['abstract'], abstract_style))
        elements.append(Spacer(1, 0.1*inch))
    
    elements.append(Paragraph(f"<b>DOI:</b> {paper.get('doi', 'N/A')}", abstract_style))
    elements.append(Spacer(1, 0.15*inch))
    
    sections_order = ['introduction', 'literature_review', 'methodology', 'results', 'discussion', 'conclusion', 'references']
    section_titles = {
        'introduction': 'I. INTRODUCTION',
        'literature_review': 'II. LITERATURE REVIEW',
        'methodology': 'III. METHODOLOGY',
        'results': 'IV. RESULTS',
        'discussion': 'V. DISCUSSION',
        'conclusion': 'VI. CONCLUSION',
        'references': 'REFERENCES'
    }
    
    for key in sections_order:
        if key in paper.get('sections', {}):
            elements.append(Paragraph(section_titles[key], section_style))
            content = paper['sections'][key]
            
            for para in content.split('\n\n'):
                if para.strip():
                    elements.append(Paragraph(para.strip(), body_style))
            
            if key == 'results' and 'figures' in paper:
                elements.append(Spacer(1, 0.1*inch))
                
                if 'figure1' in paper['figures']:
                    try:
                        fig_data = base64.b64decode(paper['figures']['figure1']['data'])
                        fig_buf = BytesIO(fig_data)
                        img = RLImage(fig_buf, width=frame_width*0.9, height=2.5*inch)
                        elements.append(img)
                        elements.append(Paragraph(paper['figures']['figure1']['caption'], caption_style))
                        elements.append(Spacer(1, 0.1*inch))
                    except:
                        pass
                
                if 'table1' in paper['figures']:
                    try:
                        table_data = paper['figures']['table1']['data']
                        t = Table(table_data, colWidths=[frame_width*0.18]*5)
                        t.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Times-Bold'),
                            ('FONTSIZE', (0, 0), (-1, -1), 7),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
                        ]))
                        elements.append(t)
                        elements.append(Paragraph(paper['figures']['table1']['caption'], caption_style))
                        elements.append(Spacer(1, 0.1*inch))
                    except:
                        pass
                
                if 'figure2' in paper['figures']:
                    try:
                        fig_data = base64.b64decode(paper['figures']['figure2']['data'])
                        fig_buf = BytesIO(fig_data)
                        img = RLImage(fig_buf, width=frame_width*0.9, height=2.5*inch)
                        elements.append(img)
                        elements.append(Paragraph(paper['figures']['figure2']['caption'], caption_style))
                    except:
                        pass
            
            elements.append(Spacer(1, 0.08*inch))
    
    doc.build(elements)
    buffer.seek(0)
    
    filename = f"IEEE_Paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/pdf')

@app.route('/api/export-docx', methods=['POST'])
def export_docx():
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        data = request.json
        paper = data.get('paper', {})
        
        if not paper:
            return jsonify({"success": False, "error": "No paper"}), 400
        
        doc = Document()
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        section = doc.sections[0]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0] if sectPr.xpath('./w:cols') else OxmlElement('w:cols')
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '288')
        if not sectPr.xpath('./w:cols'):
            sectPr.append(cols)
        
        title = doc.add_heading(paper.get('title', 'Research Paper'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.name = 'Times New Roman'
        
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(f"{paper.get('author', 'Author')}\n")
        author_run.font.size = Pt(10)
        author_run.font.name = 'Times New Roman'
        
        affil_run = author_para.add_run(f"{paper.get('affiliation', 'Institution')}\n")
        affil_run.font.size = Pt(9)
        affil_run.font.italic = True
        affil_run.font.name = 'Times New Roman'
        
        email_run = author_para.add_run(f"Email: {paper.get('email', 'N/A')}")
        email_run.font.size = Pt(9)
        email_run.font.name = 'Times New Roman'
        
        doc.add_paragraph()
        
        if 'abstract' in paper.get('sections', {}):
            abstract_heading = doc.add_paragraph()
            abstract_heading_run = abstract_heading.add_run('Abstract—')
            abstract_heading_run.font.bold = True
            abstract_heading_run.font.italic = True
            abstract_heading_run.font.size = Pt(9)
            abstract_heading_run.font.name = 'Times New Roman'
            
            abstract_para = doc.add_paragraph(paper['sections']['abstract'])
            abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in abstract_para.runs:
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.name = 'Times New Roman'
            
            doc.add_paragraph()
        
        doi_para = doc.add_paragraph()
        doi_run = doi_para.add_run(f"DOI: {paper.get('doi', 'N/A')}")
        doi_run.font.size = Pt(8)
        doi_run.font.name = 'Times New Roman'
        
        doc.add_paragraph()
        
        sections_order = ['introduction', 'literature_review', 'methodology', 'results', 'discussion', 'conclusion', 'references']
        section_titles = {
            'introduction': 'I. INTRODUCTION',
            'literature_review': 'II. LITERATURE REVIEW',
            'methodology': 'III. METHODOLOGY',
            'results': 'IV. RESULTS',
            'discussion': 'V. DISCUSSION',
            'conclusion': 'VI. CONCLUSION',
            'references': 'REFERENCES'
        }
        
        for key in sections_order:
            if key in paper.get('sections', {}):
                heading = doc.add_heading(section_titles[key], level=1)
                heading_run = heading.runs[0]
                heading_run.font.size = Pt(10)
                heading_run.font.bold = True
                heading_run.font.name = 'Times New Roman'
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                content = paper['sections'][key]
                paragraphs = content.split('\n\n')
                
                for para_text in paragraphs:
                    if para_text.strip():
                        body_para = doc.add_paragraph(para_text.strip())
                        body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        for run in body_para.runs:
                            run.font.size = Pt(9)
                            run.font.name = 'Times New Roman'
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"IEEE_Paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return send_file(buffer, as_attachment=True, download_name=filename,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    except ImportError:
        return jsonify({"success": False, "error": "Install: pip install python-docx"}), 500
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

if __name__ == '__main__':
    print("\n" + "="*70)
    print("🎓 AI Research Paper Generator - Professional Edition")
    print("="*70)
    print("📝 Ollama: ollama run mistral")
    print("🌐 URL: http://localhost:8080")
    print("⚡ Features: Fast | Plagiarism | IEEE Format | Figures | OCR")
    print("="*70 + "\n")
    app.run(debug=True, port=8080, host='0.0.0.0')
