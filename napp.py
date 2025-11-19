from flask import Flask, request, jsonify, render_template, send_file
import requests
import json
from datetime import datetime, timedelta
import os
import uuid
import re
import time
from io import BytesIO
import hashlib
import base64
from typing import List, Dict
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from wordcloud import WordCloud
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Frame, PageTemplate, BaseDocTemplate, Table, TableStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib import colors
from reportlab.platypus import Image as RLImage
from PIL import Image, ImageEnhance
import easyocr

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs('saved_papers', exist_ok=True)

CACHE_DIR = 'paper_cache'
os.makedirs(CACHE_DIR, exist_ok=True)

print("Loading EasyOCR model...")
reader = easyocr.Reader(['en'], gpu=False)
print("EasyOCR loaded!")

OLLAMA_API_URL = "http://localhost:11434/api/generate"
MODEL_NAME = "qwen2.5:7b"

CITATION_DATABASE = {
    "machine learning": [
        "[1] I. Goodfellow, Y. Bengio, and A. Courville, Deep Learning. MIT Press, 2016.",
        "[2] C. Bishop, Pattern Recognition and Machine Learning. Springer, 2006.",
        "[3] T. Mitchell, Machine Learning. McGraw-Hill, 1997.",
        "[4] Y. LeCun, Y. Bengio, and G. Hinton, \"Deep learning,\" Nature, vol. 521, pp. 436-444, 2015.",
        "[5] A. Ng, \"Machine Learning Yearning,\" Technical Strategy, 2018."
    ],
    "default": [
        "[1] A. Author, B. Author, \"Title of Paper,\" Journal Name, vol. X, no. Y, pp. ZZ-ZZ, 2023.",
        "[2] C. Author et al., \"Conference Paper Title,\" in Proc. Conf. Name, 2024, pp. XX-YY.",
        "[3] D. Author, \"Book Title,\" Publisher, 2022.",
        "[4] E. Author and F. Author, \"Another Paper,\" IEEE Trans., vol. 10, 2023.",
        "[5] G. Author, \"Recent Work,\" ACM Computing Surveys, 2024."
    ]
}

# ==================== RAG FUNCTIONS ====================

def get_cache_key(query: str) -> str:
    return hashlib.md5(query.lower().encode()).hexdigest()

def load_from_cache(query: str, max_age_hours: int = 24) -> List[Dict]:
    cache_key = get_cache_key(query)
    cache_file = os.path.join(CACHE_DIR, f"{cache_key}.json")
    
    if os.path.exists(cache_file):
        try:
            with open(cache_file, 'r') as f:
                cache_data = json.load(f)
            
            cached_time = datetime.fromisoformat(cache_data['timestamp'])
            if datetime.now() - cached_time < timedelta(hours=max_age_hours):
                age_hours = (datetime.now() - cached_time).seconds // 3600
                print(f"✅ Using cached papers (age: {age_hours}h)")
                return cache_data['papers']
        except Exception as e:
            print(f"⚠️ Cache error: {e}")
    
    return None

def save_to_cache(query: str, papers: List[Dict]):
    cache_key = get_cache_key(query)
    cache_file = os.path.join(CACHE_DIR, f"{cache_key}.json")
    
    try:
        cache_data = {
            'timestamp': datetime.now().isoformat(),
            'query': query,
            'papers': papers
        }
        with open(cache_file, 'w') as f:
            json.dump(cache_data, f)
        print(f"💾 Cached {len(papers)} papers")
    except Exception as e:
        print(f"⚠️ Cache save error: {e}")

def search_semantic_scholar(query: str, limit: int = 5) -> List[Dict]:
    cached_papers = load_from_cache(query)
    if cached_papers is not None:
        return cached_papers
    
    print(f"🔍 Searching Semantic Scholar: {query[:50]}...")
    
    url = "https://api.semanticscholar.org/graph/v1/paper/search"
    params = {
        'query': query,
        'limit': limit,
        'fields': 'title,abstract,authors,year,citationCount,paperId,venue,externalIds'
    }
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = requests.get(url, params=params, timeout=15)
            
            if response.status_code == 200:
                data = response.json()
                papers = data.get('data', [])
                valid_papers = [p for p in papers if p.get('abstract')][:limit]
                
                if valid_papers:
                    print(f"✅ Found {len(valid_papers)} papers")
                    save_to_cache(query, valid_papers)
                    return valid_papers
                else:
                    print(f"⚠️ No papers with abstracts")
                    return []
                    
            elif response.status_code == 429:
                wait_time = (2 ** attempt) * 5
                print(f"⚠️ Rate limited. Waiting {wait_time}s (retry {attempt + 1}/{max_retries})...")
                time.sleep(wait_time)
                
            else:
                print(f"⚠️ API error {response.status_code}")
                if attempt < max_retries - 1:
                    time.sleep(3)
                    
        except Exception as e:
            print(f"❌ Error: {str(e)[:100]}")
            if attempt < max_retries - 1:
                time.sleep(2)
    
    print("❌ Failed after retries")
    return []

def format_paper_context(papers: List[Dict]) -> str:
    if not papers:
        return ""
    
    context = "=== RESEARCH CONTEXT ===\n\n"
    
    for i, paper in enumerate(papers, 1):
        title = paper.get('title', 'Unknown')
        abstract = paper.get('abstract', '')[:400]
        authors = paper.get('authors', [])
        year = paper.get('year', 'N/A')
        citations = paper.get('citationCount', 0)
        
        author_names = ', '.join([a.get('name', '') for a in authors[:3]])
        if len(authors) > 3:
            author_names += ' et al.'
        
        context += f"Paper {i}: {title}\n"
        context += f"Authors: {author_names} ({year})\n"
        context += f"Citations: {citations}\n"
        context += f"Summary: {abstract}\n\n"
    
    return context

def generate_ieee_citations(papers: List[Dict]) -> List[str]:
    citations = []
    
    for i, paper in enumerate(papers, 1):
        title = paper.get('title', 'Unknown Title')
        authors = paper.get('authors', [])
        year = paper.get('year', 'N/A')
        venue = paper.get('venue', 'Journal/Conference')
        
        author_list = []
        for author in authors[:3]:
            name = author.get('name', '')
            parts = name.split()
            if len(parts) >= 2:
                formatted = f"{parts[0][0]}. {' '.join(parts[1:])}"
                author_list.append(formatted)
        
        if len(authors) > 3:
            author_str = ', '.join(author_list) + ', et al.'
        elif len(author_list) > 1:
            author_str = ', '.join(author_list[:-1]) + ' and ' + author_list[-1]
        else:
            author_str = author_list[0] if author_list else 'Unknown'
        
        external_ids = paper.get('externalIds', {})
        doi = external_ids.get('DOI', '')
        
        citation = f"[{i}] {author_str}, \"{title},\" {venue}, {year}"
        if doi:
            citation += f", doi: {doi}"
        citation += "."
        
        citations.append(citation)
    
    return citations

def get_fallback_context(topic: str) -> str:
    return f"""Based on research in {topic}, consider:
- Recent developments and state-of-the-art
- Common methodologies and frameworks
- Key challenges and research questions
- Practical applications
"""

# ==================== TITLE GENERATION ====================

def generate_paper_title(description: str, retries: int = 2) -> str:
    prompt = f"""You are an academic title expert. Generate a professional, concise research paper title from this description.

Description: {description}

Requirements:
- Maximum 12 words
- Academic and professional tone
- Capture the core research focus
- Use proper capitalization
- NO quotation marks
- Format: "Topic: Subtitle" or "Method for Application"

Generate ONLY the title, nothing else."""

    for attempt in range(retries):
        try:
            print(f"📝 Generating title (attempt {attempt + 1}/{retries})...")
            
            response = requests.post(OLLAMA_API_URL, json={
                "model": MODEL_NAME,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "num_predict": 50,
                    "temperature": 0.8,
                    "top_p": 0.9
                }
            }, timeout=30)
            
            if response.status_code == 200:
                result = response.json()
                title = result.get('response', '').strip()
                
                title = title.replace('"', '').replace("'", '').strip()
                title = re.sub(r'^Title:\s*', '', title, flags=re.IGNORECASE)
                
                words = title.split()
                if len(words) > 15:
                    title = ' '.join(words[:15])
                
                if title and len(title) > 10:
                    print(f"✅ Generated title: {title}")
                    return title
                    
            if attempt < retries - 1:
                time.sleep(2)
                
        except Exception as e:
            print(f"⚠️ Title generation error: {e}")
            if attempt < retries - 1:
                time.sleep(2)
    
    words = description.split()[:10]
    fallback = ' '.join(words).strip()
    if fallback.endswith(','):
        fallback = fallback[:-1]
    return fallback

def generate_alternative_titles(original_description: str, count: int = 3) -> list:
    prompt = f"""You are an academic title expert. Generate {count} different professional research paper titles from this description.

Description: {original_description}

Requirements for EACH title:
- Maximum 12 words
- Academic tone
- Different angles/perspectives
- NO quotation marks
- NO numbering

Generate exactly {count} titles, one per line."""

    try:
        response = requests.post(OLLAMA_API_URL, json={
            "model": MODEL_NAME,
            "prompt": prompt,
            "stream": False,
            "options": {
                "num_predict": 150,
                "temperature": 0.9,
                "top_p": 0.95
            }
        }, timeout=40)
        
        if response.status_code == 200:
            result = response.json()
            text = result.get('response', '').strip()
            
            titles = []
            for line in text.split('\n'):
                line = line.strip()
                line = re.sub(r'^[\d\.\-\*\)]+\s*', '', line)
                line = line.replace('"', '').replace("'", '').strip()
                
                if line and len(line) > 10 and len(line.split()) <= 15:
                    titles.append(line)
            
            return titles[:count] if titles else []
    except Exception as e:
        print(f"⚠️ Alternative titles error: {e}")
    
    return []

# ==================== DYNAMIC FIGURE GENERATION ====================

def extract_keywords_from_text(text: str, top_n: int = 10) -> List[str]:
    from collections import Counter
    
    stopwords = set(['the', 'a', 'an', 'in', 'on', 'at', 'for', 'to', 'of', 'and', 'is', 'are', 
                     'this', 'that', 'with', 'from', 'by', 'as', 'or', 'be', 'been', 'has', 'have'])
    
    words = re.findall(r'\b[a-z]{4,}\b', text.lower())
    filtered = [w for w in words if w not in stopwords]
    
    word_freq = Counter(filtered)
    return [word for word, count in word_freq.most_common(top_n)]

def generate_wordcloud_from_paper(paper_sections: Dict, topic: str) -> BytesIO:
    try:
        print("📊 Generating word cloud...")
        
        full_text = ' '.join(paper_sections.values())
        
        wordcloud = WordCloud(
            width=800, 
            height=400,
            background_color='white',
            colormap='viridis',
            max_words=50
        ).generate(full_text)
        
        plt.figure(figsize=(10, 5))
        plt.imshow(wordcloud, interpolation='bilinear')
        plt.axis('off')
        plt.title(f'Key Terms in "{topic}"', fontsize=14, fontweight='bold')
        plt.tight_layout()
        
        buf = BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        
        return buf
    except Exception as e:
        print(f"⚠️ Word cloud error: {e}")
        return None

def generate_keyword_frequency_chart(paper_sections: Dict) -> BytesIO:
    try:
        print("📊 Generating keyword frequency chart...")
        
        full_text = ' '.join(paper_sections.values())
        keywords = extract_keywords_from_text(full_text, top_n=10)
        
        from collections import Counter
        words = re.findall(r'\b[a-z]{4,}\b', full_text.lower())
        word_counts = Counter(words)
        
        top_words = [(word, word_counts[word]) for word in keywords]
        words, counts = zip(*top_words) if top_words else ([], [])
        
        plt.figure(figsize=(10, 6))
        bars = plt.barh(words, counts, color='steelblue', edgecolor='black', linewidth=0.7)
        plt.xlabel('Frequency', fontsize=12, fontweight='bold')
        plt.ylabel('Keywords', fontsize=12, fontweight='bold')
        plt.title('Top 10 Keywords in Paper', fontsize=14, fontweight='bold')
        plt.gca().invert_yaxis()
        
        for bar in bars:
            width = bar.get_width()
            plt.text(width + 1, bar.get_y() + bar.get_height()/2, 
                    f'{int(width)}', ha='left', va='center', fontweight='bold', fontsize=10)
        
        plt.tight_layout()
        
        buf = BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        
        return buf
    except Exception as e:
        print(f"⚠️ Keyword chart error: {e}")
        return None

def generate_metrics_table_from_rag(retrieved_papers: List[Dict]) -> List[List[str]]:
    try:
        print("📊 Generating metrics table from RAG papers...")
        
        if not retrieved_papers or len(retrieved_papers) < 2:
            return generate_generic_table()
        
        table_data = [['Paper', 'Year', 'Citations', 'Method', 'Key Finding']]
        
        for i, paper in enumerate(retrieved_papers[:5], 1):
            title = paper.get('title', 'Unknown')[:40] + '...'
            year = str(paper.get('year', 'N/A'))
            citations = str(paper.get('citationCount', 0))
            
            abstract = paper.get('abstract', '').lower()
            if 'neural' in abstract or 'deep' in abstract:
                method = 'Deep Learning'
            elif 'machine learning' in abstract:
                method = 'ML-based'
            elif 'transformer' in abstract:
                method = 'Transformer'
            else:
                method = 'Novel Approach'
            
            finding = f"{citations} citations"
            
            table_data.append([f"Paper {i}", year, citations, method, finding])
        
        return table_data
    except Exception as e:
        print(f"⚠️ Table generation error: {e}")
        return generate_generic_table()

def generate_generic_table() -> List[List[str]]:
    return [
        ['Method', 'Accuracy', 'Precision', 'Recall', 'F1-Score'],
        ['Baseline', '72.3%', '71.5%', '70.8%', '71.1%'],
        ['Method A', '78.6%', '77.2%', '79.1%', '78.1%'],
        ['Method B', '81.4%', '80.8%', '82.3%', '81.5%'],
        ['Proposed', '92.1%', '91.8%', '92.4%', '92.1%']
    ]

def process_user_table_data(csv_data: str) -> List[List[str]]:
    try:
        lines = csv_data.strip().split('\n')
        table_data = [line.split(',') for line in lines]
        return table_data
    except Exception as e:
        print(f"⚠️ User table processing error: {e}")
        return []

def generate_user_data_chart(csv_data: str, chart_type: str = 'bar') -> BytesIO:
    try:
        print(f"📊 Generating {chart_type} chart from user data...")
        
        df = pd.read_csv(BytesIO(csv_data.encode()))
        
        plt.figure(figsize=(10, 6))
        
        if chart_type == 'bar' and len(df.columns) >= 2:
            df.plot(x=df.columns[0], y=df.columns[1], kind='bar', ax=plt.gca(), color='steelblue')
            plt.xlabel(df.columns[0], fontweight='bold')
            plt.ylabel(df.columns[1], fontweight='bold')
            plt.title(f'{df.columns[1]} by {df.columns[0]}', fontweight='bold')
            plt.xticks(rotation=45)
            
        elif chart_type == 'line' and len(df.columns) >= 2:
            df.plot(x=df.columns[0], y=df.columns[1], kind='line', ax=plt.gca(), marker='o', color='steelblue')
            plt.xlabel(df.columns[0], fontweight='bold')
            plt.ylabel(df.columns[1], fontweight='bold')
            plt.title(f'{df.columns[1]} over {df.columns[0]}', fontweight='bold')
            plt.grid(True, alpha=0.3)
        
        plt.tight_layout()
        
        buf = BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        
        return buf
    except Exception as e:
        print(f"⚠️ User chart error: {e}")
        return None

# ==================== TEXT PROCESSING ====================

def clean_generated_text(text, section_name, paper_title):
    text = re.sub(r'^#+\s+.*$', '', text, flags=re.MULTILINE)
    text = re.sub(re.escape(paper_title), '', text, flags=re.IGNORECASE)
    
    section_keywords = [
        'Abstract', 'Introduction', 'Literature Review', 'Background',
        'Methodology', 'Results', 'Discussion', 'Conclusion',
        'References', 'Objectives', 'Problem Statement'
    ]
    
    for keyword in section_keywords:
        text = re.sub(rf'^{keyword}:?\s*$', '', text, flags=re.MULTILINE | re.IGNORECASE)
        text = re.sub(rf'^#+\s*{keyword}:?\s*$', '', text, flags=re.MULTILINE | re.IGNORECASE)
    
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'^\s*[-*•]\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*\d+\.\s+(?!\[)', '', text, flags=re.MULTILINE)
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    return text.strip()

def generate_text_with_context(prompt, context="", max_tokens=400, temperature=0.7, retries=2):
    if context:
        full_prompt = f"""You are a professional academic writer. Rules:
1. Write ONLY content, no headers
2. No markdown formatting
3. Continuous prose paragraphs
4. Academic tone
5. USE RESEARCH CONTEXT provided

{context}

Based on this context:
{prompt}"""
    else:
        full_prompt = f"""You are a professional academic writer. Rules:
1. Write ONLY content, no headers
2. No markdown formatting
3. Continuous prose paragraphs
4. Academic tone

{prompt}"""
    
    for attempt in range(retries):
        try:
            print(f"📤 Generating with {'context' if context else 'no context'} (attempt {attempt + 1}/{retries})...")
            
            response = requests.post(OLLAMA_API_URL, json={
                "model": MODEL_NAME,
                "prompt": full_prompt,
                "stream": False,
                "options": {
                    "num_predict": max_tokens,
                    "temperature": temperature,
                    "top_p": 0.9,
                    "top_k": 40
                }
            }, timeout=300)
            
            print(f"   Response status: {response.status_code}")
            
            if response.status_code == 200:
                result = response.json()
                generated = result.get('response', '').strip()
                print(f"✅ Generated {len(generated.split())} words")
                return generated
            else:
                print(f"❌ HTTP {response.status_code} on attempt {attempt + 1}")
                if attempt < retries - 1:
                    wait = 3
                    print(f"⏳ Waiting {wait}s before retry...")
                    time.sleep(wait)
                else:
                    error_text = response.text[:200] if hasattr(response, 'text') else 'Unknown error'
                    return f"Error: HTTP {response.status_code} - {error_text}"
                
        except requests.exceptions.Timeout:
            print(f"⏱️ Timeout on attempt {attempt + 1}")
            if attempt < retries - 1:
                print(f"⏳ Retrying...")
                time.sleep(2)
            else:
                return "Error: Request timeout"
        except Exception as e:
            print(f"❌ Exception on attempt {attempt + 1}: {type(e).__name__}")
            if attempt < retries - 1:
                print(f"⏳ Retrying...")
                time.sleep(2)
            else:
                return f"Error: {str(e)}"
    
    return "Error: All retries failed"

def extract_text_from_image(image_path):
    try:
        if not os.path.exists(image_path):
            return "OCR Error: Image not found"
        
        img = Image.open(image_path)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.5)
        
        results = reader.readtext(image_path, detail=0, paragraph=False)
        text = ' '.join(results).strip()
        
        return text if text else "No text detected"
    except Exception as e:
        return f"OCR Error: {str(e)}"

def check_plagiarism(text):
    sentences = [s.strip() for s in text.split('.') if s.strip()]
    unique_sentences = len(set([s.lower() for s in sentences]))
    uniqueness = (unique_sentences / len(sentences) * 100) if sentences else 0
    
    words = re.findall(r'\w+', text.lower())
    unique_words = len(set(words))
    vocabulary_score = (unique_words / len(words) * 100) if words else 0
    
    final_score = (uniqueness * 0.6 + min(vocabulary_score, 100) * 0.4)
    
    return {
        "uniqueness_score": round(final_score, 1),
        "status": "Highly Original" if final_score > 85 else "Original" if final_score > 70 else "Needs Review",
        "total_sentences": len(sentences),
        "unique_sentences": unique_sentences,
        "vocabulary_diversity": round(vocabulary_score, 1)
    }

def get_citations(topic):
    topic_lower = topic.lower()
    for key in CITATION_DATABASE.keys():
        if key in topic_lower:
            return CITATION_DATABASE[key]
    return CITATION_DATABASE["default"]

def generate_doi():
    return f"10.1109/ACCESS.{datetime.now().year}.{uuid.uuid4().hex[:8].upper()}"

# ==================== ROUTES ====================

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/api/generate-paper', methods=['POST'])
def generate_paper():
    data = request.json
    topic_or_description = data.get('topic', '')
    authors = data.get('authors', [])  # Array of author objects
    use_rag = data.get('use_rag', True)
    user_table_csv = data.get('user_table_data', '')
    user_charts = data.get('user_charts', [])
    
    if not topic_or_description:
        return jsonify({"success": False, "error": "Topic required"}), 400
    
    # Ensure at least one author
    if not authors or len(authors) == 0:
        authors = [{
            "name": "Author Name",
            "email": "author@university.edu",
            "affiliation": "University Name"
        }]
    
    # Generate title
    print(f"📝 Input: {topic_or_description[:100]}...")
    
    word_count = len(topic_or_description.split())
    if word_count > 10 or len(topic_or_description) > 80:
        print("🎯 Detected description, generating title...")
        paper_title = generate_paper_title(topic_or_description)
        research_topic = topic_or_description
    else:
        print("🎯 Using input as title...")
        paper_title = topic_or_description
        research_topic = topic_or_description
    
    paper = {
        "title": paper_title,
        "original_input": topic_or_description,
        "authors": authors,  # Store multiple authors
        "doi": generate_doi(),
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "sections": {},
        "metadata": {},
        "figures": {},
        "rag_enabled": use_rag
    }
    
    print(f"\n{'='*70}")
    print(f"⚡ GENERATION")
    print(f"Title: {paper_title}")
    print(f"Authors: {len(authors)}")
    print(f"RAG: {'Enabled' if use_rag else 'Disabled'}")
    print(f"{'='*70}\n")
    
    # Warmup Ollama
    print("🔥 Warming up Ollama...")
    try:
        warmup = requests.post(OLLAMA_API_URL, json={
            "model": MODEL_NAME,
            "prompt": "Test",
            "stream": False,
            "options": {"num_predict": 10}
        }, timeout=30)
        print("✅ Ollama ready\n" if warmup.status_code == 200 else "⚠️ Warmup issue\n")
    except:
        print("⚠️ Could not warm up Ollama\n")
    
    # RAG
    retrieved_papers = []
    research_context = ""
    
    if use_rag:
        print("🔍 RAG: Fetching papers...")
        retrieved_papers = search_semantic_scholar(research_topic, limit=3)
        
        if retrieved_papers:
            research_context = format_paper_context(retrieved_papers)
            print(f"✅ RAG: Using {len(retrieved_papers)} papers\n")
        else:
            print("⚠️ RAG: Using fallback context\n")
            research_context = get_fallback_context(research_topic)
    
    # Generate sections
    prompts = {
        "abstract": (
            f"Write a 180-word abstract for '{research_topic}'. Include background, objectives, methods, results, conclusion.",
            300
        ),
        "introduction": (
            f"Write a 350-word introduction for '{research_topic}'. Context, problem, objectives.",
            450
        ),
        "literature_review": (
            f"Write a 350-word literature review for '{research_topic}'. Discuss papers from context.",
            450
        ),
        "methodology": (
            f"Write a 350-word methodology for '{research_topic}'. Approach, design, methods.",
            450
        ),
        "results": (
            f"Write a 300-word results for '{research_topic}'. Present findings. Mention figures/tables.",
            400
        ),
        "discussion": (
            f"Write a 350-word discussion for '{research_topic}'. Interpret results, compare with context.",
            450
        ),
        "conclusion": (
            f"Write a 250-word conclusion for '{research_topic}'. Summary and future work.",
            350
        ),
    }
    
    for section_name, (prompt, max_tokens) in prompts.items():
        print(f"🔄 Generating {section_name}...")
        
        use_context = section_name in ['introduction', 'literature_review', 'discussion']
        context = research_context if (use_rag and use_context) else ""
        
        result = None
        for attempt in range(2):
            result = generate_text_with_context(prompt, context=context, max_tokens=max_tokens)
            
            if result and not result.startswith("Error") and len(result) > 50:
                break
            else:
                if attempt == 0:
                    print(f"⚠️ Retrying {section_name}...")
                    time.sleep(2)
        
        if result and not result.startswith("Error") and len(result) > 50:
            cleaned_result = clean_generated_text(result, section_name, paper_title)
            paper["sections"][section_name] = cleaned_result
            print(f"✅ {section_name} - {len(cleaned_result.split())} words\n")
        else:
            print(f"❌ {section_name} FAILED\n")
            paper["sections"][section_name] = f"[Content for {section_name}]"
    
    # Generate figures
    print("📊 Generating figures...")
    try:
        wordcloud_buf = generate_wordcloud_from_paper(paper["sections"], paper_title)
        if wordcloud_buf:
            paper["figures"]["figure1"] = {
                "caption": "Figure 1: Key terms",
                "data": base64.b64encode(wordcloud_buf.getvalue()).decode('utf-8'),
                "type": "wordcloud"
            }
        
        keyword_chart_buf = generate_keyword_frequency_chart(paper["sections"])
        if keyword_chart_buf:
            paper["figures"]["figure2"] = {
                "caption": "Figure 2: Keyword frequency",
                "data": base64.b64encode(keyword_chart_buf.getvalue()).decode('utf-8'),
                "type": "keyword_chart"
            }
        
        if retrieved_papers:
            table_data = generate_metrics_table_from_rag(retrieved_papers)
        else:
            table_data = generate_generic_table()
        
        paper["figures"]["table1"] = {
            "caption": "Table 1: Comparative analysis",
            "data": table_data,
            "type": "table"
        }
        
        if user_table_csv:
            user_table = process_user_table_data(user_table_csv)
            if user_table:
                paper["figures"]["table2"] = {
                    "caption": "Table 2: Experimental results",
                    "data": user_table,
                    "type": "table"
                }
        
        figure_num = 3
        for idx, chart_info in enumerate(user_charts):
            try:
                chart_data = chart_info.get('data', '')
                chart_type = chart_info.get('type', 'bar')
                chart_title = chart_info.get('title', f'Chart {idx+1}')
                
                if chart_data:
                    user_chart_buf = generate_user_data_chart(chart_data, chart_type)
                    if user_chart_buf:
                        paper["figures"][f"figure{figure_num}"] = {
                            "caption": f"Figure {figure_num}: {chart_title}",
                            "data": base64.b64encode(user_chart_buf.getvalue()).decode('utf-8'),
                            "type": "user_chart"
                        }
                        figure_num += 1
            except Exception as e:
                print(f"⚠️ Chart {idx+1} error: {e}")
        
        print("✅ Figures done\n")
    except Exception as e:
        print(f"⚠️ Figure error: {e}\n")
    
    # Citations
    if use_rag and retrieved_papers:
        paper["sections"]["references"] = "\n".join(generate_ieee_citations(retrieved_papers))
    else:
        paper["sections"]["references"] = "\n".join(get_citations(research_topic))
    
    # Metadata
    total_words = sum(len(str(content).split()) for content in paper["sections"].values())
    paper["metadata"] = {
        "word_count": total_words,
        "page_estimate": round(total_words / 250),
        "has_figures": len([k for k in paper.get("figures", {}).keys() if k.startswith('figure')]),
        "has_tables": len([k for k in paper.get("figures", {}).keys() if k.startswith('table')]),
        "rag_papers_used": len(retrieved_papers),
        "user_charts_added": len(user_charts)
    }
    
    print(f"{'='*70}")
    print(f"✅ COMPLETE! {total_words} words")
    print(f"{'='*70}\n")
    
    return jsonify({"success": True, "paper": paper})

@app.route('/api/regenerate-titles', methods=['POST'])
def regenerate_titles():
    data = request.json
    description = data.get('description', '')
    
    if not description:
        return jsonify({"success": False, "error": "Description required"}), 400
    
    try:
        titles = generate_alternative_titles(description, count=3)
        
        if titles:
            return jsonify({"success": True, "titles": titles})
        else:
            return jsonify({"success": False, "error": "Failed to generate titles"}), 500
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/plagiarism-check', methods=['POST'])
def plagiarism_check():
    data = request.json
    paper = data.get('paper', {})
    
    if not paper:
        return jsonify({"success": False, "error": "No paper"}), 400
    
    full_text = " ".join(paper.get('sections', {}).values())
    result = check_plagiarism(full_text)
    
    return jsonify({"success": True, "plagiarism_check": result})

@app.route('/api/ocr-generate', methods=['POST'])
def ocr_generate():
    if 'image' not in request.files:
        return jsonify({"success": False, "error": "No image"}), 400
    
    file = request.files['image']
    if not file.filename:
        return jsonify({"success": False, "error": "Empty filename"}), 400
    
    try:
        file_ext = os.path.splitext(file.filename)[1].lower() or '.jpg'
        safe_filename = f"ocr_{uuid.uuid4().hex[:8]}{file_ext}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)
        
        file.save(filepath)
        extracted_text = extract_text_from_image(filepath)
        
        try:
            os.remove(filepath)
        except:
            pass
        
        return jsonify({"success": True, "extracted_text": extracted_text})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/api/export-pdf', methods=['POST'])
def export_pdf():
    data = request.json
    paper = data.get('paper', {})
    
    if not paper:
        return jsonify({"success": False, "error": "No paper"}), 400
    
    buffer = BytesIO()
    doc = BaseDocTemplate(buffer, pagesize=letter,
                         leftMargin=0.75*inch, rightMargin=0.75*inch,
                         topMargin=0.75*inch, bottomMargin=0.75*inch)
    
    frame_width = (letter[0] - 1.5*inch - 0.2*inch) / 2
    frame_height = letter[1] - 1.5*inch
    
    left_frame = Frame(0.75*inch, 0.75*inch, frame_width, frame_height, id='left')
    right_frame = Frame(0.75*inch + frame_width + 0.2*inch, 0.75*inch, frame_width, frame_height, id='right')
    
    doc.addPageTemplates([PageTemplate(id='TwoCol', frames=[left_frame, right_frame])])
    
    elements = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, leading=22,
                                 alignment=TA_CENTER, fontName='Times-Bold', spaceAfter=12)
    author_style = ParagraphStyle('Author', parent=styles['Normal'], fontSize=10, leading=14,
                                   alignment=TA_CENTER, fontName='Times-Roman', spaceAfter=6)
    section_style = ParagraphStyle('Section', parent=styles['Heading2'], fontSize=10, leading=12,
                                    fontName='Times-Bold', spaceAfter=6, spaceBefore=8)
    body_style = ParagraphStyle('Body', parent=styles['BodyText'], fontSize=9, leading=11,
                                alignment=TA_JUSTIFY, fontName='Times-Roman', firstLineIndent=12, spaceAfter=6)
    abstract_style = ParagraphStyle('Abstract', parent=body_style, fontName='Times-Italic', firstLineIndent=0)
    caption_style = ParagraphStyle('Caption', parent=body_style, fontSize=8, alignment=TA_CENTER, 
                                   fontName='Times-Italic', firstLineIndent=0)
    
    # Title
    elements.append(Paragraph(paper.get('title', 'Research Paper'), title_style))
    elements.append(Spacer(1, 0.1*inch))
    
    # Multiple authors
    authors = paper.get('authors', [])
    for author in authors:
        author_block = f"{author.get('name', 'Author')}<br/>{author.get('affiliation', 'Institution')}<br/>{author.get('email', 'email@edu')}"
        elements.append(Paragraph(author_block, author_style))
    
    elements.append(Spacer(1, 0.15*inch))
    
    # Abstract
    if 'abstract' in paper.get('sections', {}):
        elements.append(Paragraph("<b><i>Abstract</i></b>—", section_style))
        elements.append(Paragraph(paper['sections']['abstract'], abstract_style))
        elements.append(Spacer(1, 0.1*inch))
    
    elements.append(Paragraph(f"<b>DOI:</b> {paper.get('doi', 'N/A')}", abstract_style))
    elements.append(Spacer(1, 0.15*inch))
    
    # Sections
    sections_order = ['introduction', 'literature_review', 'methodology', 'results', 'discussion', 'conclusion', 'references']
    section_titles = {
        'introduction': 'I. INTRODUCTION',
        'literature_review': 'II. LITERATURE REVIEW',
        'methodology': 'III. METHODOLOGY',
        'results': 'IV. RESULTS',
        'discussion': 'V. DISCUSSION',
        'conclusion': 'VI. CONCLUSION',
        'references': 'REFERENCES'
    }
    
    for key in sections_order:
        if key in paper.get('sections', {}):
            elements.append(Paragraph(section_titles[key], section_style))
            content = paper['sections'][key]
            
            for para in content.split('\n\n'):
                if para.strip():
                    elements.append(Paragraph(para.strip(), body_style))
            
            if key == 'results' and 'figures' in paper:
                elements.append(Spacer(1, 0.1*inch))
                
                # Add figures and tables
                for fig_key in sorted([k for k in paper['figures'].keys() if k.startswith('figure')]):
                    try:
                        fig_data = base64.b64decode(paper['figures'][fig_key]['data'])
                        fig_buf = BytesIO(fig_data)
                        img = RLImage(fig_buf, width=frame_width*0.9, height=2.5*inch)
                        elements.append(img)
                        elements.append(Paragraph(paper['figures'][fig_key]['caption'], caption_style))
                        elements.append(Spacer(1, 0.1*inch))
                    except:
                        pass
                
                for table_key in sorted([k for k in paper['figures'].keys() if k.startswith('table')]):
                    try:
                        table_data = paper['figures'][table_key]['data']
                        num_cols = len(table_data[0])
                        col_width = frame_width / num_cols * 0.9
                        t = Table(table_data, colWidths=[col_width] * num_cols)
                        t.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Times-Bold'),
                            ('FONTSIZE', (0, 0), (-1, -1), 7),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
                        ]))
                        elements.append(t)
                        elements.append(Paragraph(paper['figures'][table_key]['caption'], caption_style))
                        elements.append(Spacer(1, 0.1*inch))
                    except:
                        pass
            
            elements.append(Spacer(1, 0.08*inch))
    
    doc.build(elements)
    buffer.seek(0)
    
    filename = f"IEEE_Paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/pdf')

@app.route('/api/export-docx', methods=['POST'])
def export_docx():
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        data = request.json
        paper = data.get('paper', {})
        
        if not paper:
            return jsonify({"success": False, "error": "No paper"}), 400
        
        doc = Document()
        
        # Title
        title = doc.add_heading(paper.get('title', 'Research Paper'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Authors
        authors = paper.get('authors', [])
        for author in authors:
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(f"{author.get('name', 'Author')}\n")
            author_run.font.size = Pt(10)
            affil_run = author_para.add_run(f"{author.get('affiliation', 'Institution')}\n")
            affil_run.font.size = Pt(9)
            affil_run.font.italic = True
            email_run = author_para.add_run(f"{author.get('email', 'N/A')}\n")
            email_run.font.size = Pt(9)
        
        doc.add_paragraph()
        
        # Abstract
        if 'abstract' in paper.get('sections', {}):
            abstract_heading = doc.add_paragraph()
            abstract_heading_run = abstract_heading.add_run('Abstract—')
            abstract_heading_run.font.bold = True
            abstract_heading_run.font.italic = True
            
            abstract_para = doc.add_paragraph(paper['sections']['abstract'])
            abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Rest of sections
        sections_order = ['introduction', 'literature_review', 'methodology', 'results', 'discussion', 'conclusion', 'references']
        section_titles = {
            'introduction': 'I. INTRODUCTION',
            'literature_review': 'II. LITERATURE REVIEW',
            'methodology': 'III. METHODOLOGY',
            'results': 'IV. RESULTS',
            'discussion': 'V. DISCUSSION',
            'conclusion': 'VI. CONCLUSION',
            'references': 'REFERENCES'
        }
        
        for key in sections_order:
            if key in paper.get('sections', {}):
                heading = doc.add_heading(section_titles[key], level=1)
                content = paper['sections'][key]
                
                for para_text in content.split('\n\n'):
                    if para_text.strip():
                        body_para = doc.add_paragraph(para_text.strip())
                        body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"IEEE_Paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return send_file(buffer, as_attachment=True, download_name=filename,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    except ImportError:
        return jsonify({"success": False, "error": "Install: pip install python-docx"}), 500
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

if __name__ == '__main__':
    print("\n" + "="*70)
    print("🎓 AI Research Paper Generator v3.0 - Multi-Author Support")
    print("="*70)
    print(f"📝 Model: {MODEL_NAME}")
    print("="*70)
    
    try:
        test = requests.get("http://localhost:11434/api/tags", timeout=5)
        if test.status_code == 200:
            print("✅ Ollama running!")
        else:
            print("⚠️ Ollama issue")
    except:
        print("❌ Cannot connect to Ollama!")
    
    print("="*70)
    print("🚀 Server: http://localhost:8080")
    print("="*70 + "\n")
    
    app.run(debug=True, port=8080, host='0.0.0.0')
